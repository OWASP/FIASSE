"""Convert markdown documents to DOCX.

Default behavior converts docs/securable_framework.md to docs/securable_framework.docx.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document


HEADING_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6"}


def log_status(level: str, message: str, **context: object) -> None:
    """Print structured status output with optional key=value context."""
    if context:
        context_str = " ".join(f"{key}={value}" for key, value in sorted(context.items()))
        print(f"[{level}] {message} | {context_str}")
        return
    print(f"[{level}] {message}")


def normalize_whitespace(text: str) -> str:
    """Normalize runs of whitespace while preserving single spaces between words."""
    return re.sub(r"\s+", " ", text)


def add_inline_content(paragraph, node: Tag | NavigableString, code_font_name: str = "Consolas") -> None:
    """Recursively append inline HTML content into a DOCX paragraph."""
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(normalize_whitespace(text))
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    if node.name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = code_font_name
        return

    if node.name in {"strong", "b", "em", "i", "a", "span"}:
        start_runs = len(paragraph.runs)
        for child in node.children:
            add_inline_content(paragraph, child, code_font_name)
        for run in paragraph.runs[start_runs:]:
            if node.name in {"strong", "b"}:
                run.bold = True
            if node.name in {"em", "i"}:
                run.italic = True
        if node.name == "a" and node.has_attr("href"):
            href = node["href"]
            for run in paragraph.runs[start_runs:]:
                run.underline = True
            paragraph.add_run(f" ({href})")
        return

    for child in node.children:
        add_inline_content(paragraph, child, code_font_name)


def add_list(document: Document, list_tag: Tag, level: int = 0) -> None:
    """Render UL/OL lists with simple nested support."""
    style_name = "List Bullet" if list_tag.name == "ul" else "List Number"
    indent_prefix = "  " * level

    for li in list_tag.find_all("li", recursive=False):
        para = document.add_paragraph(style=style_name)
        if level:
            para.add_run(indent_prefix)

        for child in li.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                add_list(document, child, level + 1)
            else:
                add_inline_content(para, child)


def add_blockquote(document: Document, quote_tag: Tag) -> None:
    """Render blockquotes as paragraphs prefixed by a quote marker."""
    text = normalize_whitespace(quote_tag.get_text(" ", strip=True))
    if text:
        document.add_paragraph(f"> {text}")


def add_code_block(document: Document, pre_tag: Tag) -> None:
    """Render fenced/indented code blocks using monospaced font."""
    code_text = pre_tag.get_text("\n")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code_text.rstrip("\n"))
    run.font.name = "Consolas"


def html_to_docx_document(html: str) -> Document:
    """Convert a subset of HTML to a DOCX Document."""
    document = Document()
    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup

    for element in root.children:
        if not isinstance(element, Tag):
            continue

        if element.name in HEADING_TAGS:
            level = min(int(element.name[1]), 9)
            text = normalize_whitespace(element.get_text(" ", strip=True))
            document.add_heading(text, level=level)
        elif element.name == "p":
            paragraph = document.add_paragraph()
            for child in element.children:
                add_inline_content(paragraph, child)
        elif element.name in {"ul", "ol"}:
            add_list(document, element)
        elif element.name == "blockquote":
            add_blockquote(document, element)
        elif element.name == "pre":
            add_code_block(document, element)
        elif element.name == "hr":
            document.add_paragraph("-" * 20)
        elif element.name == "table":
            # Keep a readable plain-text representation for table content.
            for row in element.find_all("tr"):
                cells = [normalize_whitespace(cell.get_text(" ", strip=True)) for cell in row.find_all(["th", "td"])]
                if cells:
                    document.add_paragraph(" | ".join(cells))
        else:
            text = normalize_whitespace(element.get_text(" ", strip=True))
            if text:
                document.add_paragraph(text)

    return document


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    """Convert a markdown file to DOCX."""
    markdown_text = input_path.read_text(encoding="utf-8")
    html = markdown.markdown(
        markdown_text,
        extensions=[
            "extra",
            "tables",
            "fenced_code",
            "sane_lists",
            "toc",
        ],
        output_format="html5",
    )

    document = html_to_docx_document(html)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(description="Convert markdown files to DOCX.")
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("docs/securable_framework.md"),
        help="Path to input markdown file (default: docs/securable_framework.md)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/securable_framework.docx"),
        help="Path to output docx file (default: docs/securable_framework.docx)",
    )
    return parser.parse_args()


def main() -> int:
    """CLI entrypoint."""
    args = parse_args()
    input_path = args.input.resolve()
    output_path = args.output.resolve()

    if not input_path.exists():
        log_status("ERROR", "Input markdown file does not exist", path=input_path)
        return 1

    try:
        convert_markdown_to_docx(input_path, output_path)
    except Exception as exc:  # pragma: no cover - defensive error boundary
        log_status("ERROR", "Failed to convert markdown to DOCX", reason=str(exc))
        return 1

    log_status("INFO", "DOCX generated", input=input_path, output=output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
